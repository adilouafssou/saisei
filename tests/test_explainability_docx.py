"""Verifier for the explainability-report DOCX exporter (Feature 7, Word path).

The Markdown report + its determinism/numeric invariants are covered by
``tests/test_explainability_report.py``. This pins the SEPARATE concern of the
Word (.docx) artifact built from that SAME report:

1. **It is a real .docx.** Non-empty bytes beginning with the ZIP ``PK``
   signature, openable by ``python-docx``.
2. **Numeric preservation.** Converting the report to Word adds, drops, or
   alters NO figure: the multiset of yen values AND the EWS / pillar scores in
   the .docx (paragraphs + table cells) equals the Markdown report's.
3. **Content + structure carried.** The borrower name, FSA class kanji, and the
   classification reason survive, and the EWS breakdown Markdown table becomes a
   real Word table (not raw pipe-delimited text).
4. **Determinism.** Same state in -> same extracted text out.
5. **Filename contract.** Mirrors the Markdown/DOCX cross-platform rule with a
   ``.docx`` extension.

All tests are offline, deterministic, and import only from ``app.*`` (+ docx).
"""

from __future__ import annotations

import datetime as dt
import io
from collections import Counter

from app.backend.analysis.numeric_preservation import extract_yen_values
from app.backend.export.explainability_report import (
    build_explainability_docx,
    build_explainability_report,
    explainability_docx_filename,
)
from app.backend.nodes.ews_scoring import compute_ews_breakdown, compute_ews_score
from app.backend.nodes.keieisha_hosho import assess_hosho_kaijo
from app.backend.state import SaiseiState
from app.backend.tools.tdb_api import CompanyProfile
from app.shared.models.accounting import TrialBalance
from app.shared.models.classification import FsaClass
from docx import Document


def _profile() -> CompanyProfile:
    return CompanyProfile(
        tdb_code="1234567",
        hojin_bango="1234567890123",
        name="\u30c6\u30b9\u30c8\u88fd\u9020\u682a\u5f0f\u4f1a\u793e",  # テスト製造株式会社
        prefecture="\u611b\u77e5\u770c",
        industry="\u88fd\u9020\u696d",
        established_year=1990,
        employees=42,
    )


def _declining_history() -> list[TrialBalance]:
    rows: list[TrialBalance] = []
    for i in range(6):
        sales = 100_000_000 - i * 8_000_000
        cogs = int(sales * (0.72 + i * 0.01))
        rows.append(
            TrialBalance(
                period=dt.date(2025, 1, 1) + dt.timedelta(days=30 * i),
                uriage=sales,
                uriage_genka=cogs,
                hanbaihi=20_000_000,
            )
        )
    return rows


def _yochuisaki_state() -> SaiseiState:
    """A fully-assessed 要注意先 borrower (with EWS breakdown + Hosho pillars)."""
    history = _declining_history()
    conditions = assess_hosho_kaijo(
        shisanhyo_count=len(history),
        avg_eigai_shueki=0.0,
        avg_eigai_hiyo=0.0,
        avg_uriage=90_000_000.0,
        ews_score=compute_ews_score(history),
        working_capital_gap=-5_000_000,
        tdb_score=58,
        error_count=0,
    )
    return SaiseiState(
        tdb_code="1234567",
        company_profile=_profile(),
        tdb_score=58,
        shisanhyo=history,
        working_capital_gap=-5_000_000,
        ews_score=compute_ews_score(history),
        ews_breakdown=[s.__dict__ for s in compute_ews_breakdown(history)],
        fsa_classification=FsaClass.YOCHUISAKI,
        special_attention=True,
        classification_reason=(
            "\u8cc7\u91d1\u7e70\u308a\u4e0d\u8db3 "
            "(working-capital deficit \u2192 \u8981\u7ba1\u7406\u5148)"
        ),
        hosho_kaijo_score=round(
            conditions.bunri_score + conditions.zaimu_score + conditions.kaiji_score, 2
        ),
        hosho_kaijo_conditions=conditions,
    )


def _docx_all_text(data: bytes) -> str:
    """Extract paragraph AND table-cell text from generated ``.docx`` bytes.

    ``python-docx`` exposes table cells separately from ``document.paragraphs``,
    so a check that ignored tables would miss any figure moved into a Word table.
    This walks both, exactly like ``tests/test_keikakusho_docx.py``.
    """
    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_docx_is_a_real_word_file() -> None:
    """Non-empty bytes with the ZIP (PK) signature, openable by python-docx."""
    data = build_explainability_docx(_yochuisaki_state())
    assert isinstance(data, bytes)
    assert len(data) > 0
    assert data[:2] == b"PK"  # a .docx is a zip archive
    Document(io.BytesIO(data))  # opens without raising


def test_docx_preserves_every_yen_figure() -> None:
    """The DOCX carries exactly the same multiset of yen values as the report."""
    state = _yochuisaki_state()
    md = build_explainability_report(state)
    data = build_explainability_docx(state)
    source = Counter(extract_yen_values(md))
    rendered = Counter(extract_yen_values(_docx_all_text(data)))
    assert rendered == source


def test_docx_preserves_ews_and_pillar_scores() -> None:
    """The EWS score and each per-signal point value survive into the .docx."""
    state = _yochuisaki_state()
    text = _docx_all_text(build_explainability_docx(state))

    score = compute_ews_score(state.shisanhyo)
    score_txt = str(int(score)) if score == int(score) else f"{score:.2f}"
    assert score_txt in text

    for sig in compute_ews_breakdown(state.shisanhyo):
        pts = str(int(sig.points)) if sig.points == int(sig.points) else f"{sig.points:.2f}"
        assert pts in text


def test_docx_carries_company_and_classification() -> None:
    """The borrower name, FSA class kanji, and the threshold reason survive."""
    text = _docx_all_text(build_explainability_docx(_yochuisaki_state()))
    assert "\u30c6\u30b9\u30c8\u88fd\u9020\u682a\u5f0f\u4f1a\u793e" in text  # company name
    assert FsaClass.YOCHUISAKI.kanji in text  # 要注意先
    assert "working-capital deficit" in text  # classification reason verbatim


def test_ews_breakdown_becomes_a_word_table() -> None:
    """The EWS breakdown Markdown table renders as a real Word table."""
    document = Document(io.BytesIO(build_explainability_docx(_yochuisaki_state())))
    assert len(document.tables) >= 1
    # No raw pipe-delimited table row leaked as a paragraph, and the layout
    # delimiter row never appears as text.
    para_text = "\n".join(p.text for p in document.paragraphs)
    assert "| ---" not in para_text
    assert "---: |" not in para_text


def test_docx_is_deterministic_text() -> None:
    """Same state in -> same extracted text out (no clock, LLM, or network)."""
    state = _yochuisaki_state()
    first = _docx_all_text(build_explainability_docx(state))
    second = _docx_all_text(build_explainability_docx(state))
    assert first == second


def test_docx_filename_is_safe() -> None:
    """The .docx filename mirrors the cross-platform contract with a .docx ext."""
    assert (
        explainability_docx_filename("\u30c6\u30b9\u30c8 \u88fd\u9020")
        == "explainability_\u30c6\u30b9\u30c8_\u88fd\u9020.docx"
    )
    assert explainability_docx_filename("bad:name?") == "explainability_bad_name.docx"
    assert explainability_docx_filename("") == "explainability_borrower.docx"
