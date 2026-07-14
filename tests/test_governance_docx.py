"""Verifier for the governance-document DOCX exporters (Feature 7, Word path).

The Markdown card / change log + their determinism / no-drift invariants are
covered by ``tests/test_model_card.py``. This pins the SEPARATE concern of the
Word (.docx) artifacts built from those SAME documents:

1. **They are real .docx files.** Non-empty bytes beginning with the ZIP ``PK``
   signature, openable by ``python-docx``.
2. **No drift from the code.** Every governing constant's LIVE value survives
   into the model-card .docx (paragraphs + table cells), so the Word artifact
   cannot silently diverge from the running engine.
3. **The constants table becomes a real Word table.**
4. **Faithful change log.** A changed threshold renders old AND new in the .docx.
5. **Determinism + filename contract.**

All tests are offline, deterministic, and import only from ``app.*`` (+ docx).
"""

from __future__ import annotations

import io

from app.backend.export.model_card import (
    build_constants_changelog_docx,
    build_model_card_docx,
    constants_changelog_docx_filename,
    governing_constants,
    model_card_docx_filename,
)
from app.shared import constants as C
from docx import Document


def _fmt(value: object) -> str:
    """Mirror the module's compact int/float formatting for assertions."""
    if isinstance(value, bool):
        return str(value)
    if isinstance(value, float) and value == int(value):
        return str(int(value))
    return str(value)


def _docx_all_text(data: bytes) -> str:
    """Extract paragraph AND table-cell text from generated ``.docx`` bytes."""
    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_model_card_docx_is_a_real_word_file() -> None:
    """Non-empty bytes with the ZIP (PK) signature, openable by python-docx."""
    data = build_model_card_docx()
    assert isinstance(data, bytes)
    assert len(data) > 0
    assert data[:2] == b"PK"
    Document(io.BytesIO(data))  # opens without raising


def test_model_card_docx_carries_every_governing_constant() -> None:
    """Every governing constant's live name AND value survives into the .docx.

    The anti-drift guarantee in Word form: the card .docx is rendered from
    ``governing_constants()``, so a value change is reflected verbatim and a
    dropped constant is caught.
    """
    text = _docx_all_text(build_model_card_docx())
    for name, value in governing_constants().items():
        assert name in text, f"missing constant name: {name}"
        assert _fmt(value) in text, f"missing value for {name}: {value}"


def test_model_card_docx_constants_table_is_a_word_table() -> None:
    """The governing-constants Markdown table renders as a real Word table."""
    document = Document(io.BytesIO(build_model_card_docx()))
    assert len(document.tables) >= 1
    para_text = "\n".join(p.text for p in document.paragraphs)
    # No raw pipe-delimited row or layout delimiter leaked as a paragraph.
    assert "| ---" not in para_text


def test_model_card_docx_is_deterministic_text() -> None:
    """Same engine config in -> same extracted text out."""
    assert _docx_all_text(build_model_card_docx()) == _docx_all_text(build_model_card_docx())


def test_changelog_docx_is_a_real_word_file() -> None:
    """The change-log .docx is a real, openable Word file (first issuance)."""
    data = build_constants_changelog_docx(previous=None)
    assert data[:2] == b"PK"
    Document(io.BytesIO(data))


def test_changelog_docx_reports_changed_value_old_to_new() -> None:
    """A changed threshold renders old AND new values into the .docx."""
    baseline = dict(governing_constants())
    baseline["EWS_SUBSTANDARD"] = 35.0  # pretend the floor was previously 35
    text = _docx_all_text(build_constants_changelog_docx(previous=baseline))
    assert "EWS_SUBSTANDARD" in text
    assert "35" in text  # old
    assert _fmt(C.EWS_SUBSTANDARD) in text  # new (live)


def test_changelog_docx_is_deterministic_text() -> None:
    """Same inputs -> same extracted change-log text out."""
    baseline = dict(governing_constants())
    baseline["EWS_DANGER"] = 80.0
    first = _docx_all_text(build_constants_changelog_docx(previous=baseline))
    second = _docx_all_text(build_constants_changelog_docx(previous=baseline))
    assert first == second


def test_governance_docx_filenames_are_safe() -> None:
    """The .docx filenames mirror the cross-platform contract."""
    assert model_card_docx_filename() == "model_card_saisei_engine.docx"
    assert model_card_docx_filename("bad:name?") == "model_card_bad_name.docx"
    assert model_card_docx_filename("") == "model_card_engine.docx"
    assert constants_changelog_docx_filename() == "governing_constants_changelog.docx"
