"""Verifier for Keikakusho export delivery into a bank house DOCX template.

Proves, fully offline and deterministically (python-docx only), that injecting
the deterministic Keikakusho body into a bank template:
* preserves the template's surrounding content (cover page / fixed sections);
* removes the {{KEIKAKUSHO_BODY}} placeholder and places the body there;
* preserves every yen figure byte-for-byte (the inviolable numeric invariant);
* renders Markdown tables as real Word tables;
* raises TemplateError when the placeholder is absent;
* leaves the default (no-template) export behaviour unchanged;
* and that the settings-aware entry point falls back to the default when no
  template path is configured.
"""

from __future__ import annotations

import io

import pytest
from app.backend.export.keikakusho_docx import (
    build_keikakusho_docx,
    build_keikakusho_docx_for_settings,
)
from app.backend.export.keikakusho_template import (
    DEFAULT_BODY_PLACEHOLDER,
    TemplateError,
    build_keikakusho_docx_from_template,
)
from docx import Document

_DRAFT = """# \u7d4c\u55b6\u6539\u5584\u8a08\u753b\u66f8

## \u8cb7\u53ce\u652f\u6539\u5584

\u58f2\u4e0a\u9ad8\u76ee\u6a19\u306f \u00a5123,456,789 \u3067\u3059\u3002

| \u9805\u76ee | \u91d1\u984d |
| --- | --- |
| \u58f2\u4e0a | \u00a5100,000,000 |
| \u7d4c\u5e38\u5229\u76ca | -\u00a55,000,000 |
"""

# The exact yen figures that must survive byte-for-byte.
_FIGURES = ["\u00a5123,456,789", "\u00a5100,000,000", "-\u00a55,000,000"]


def _template_bytes(placeholder: str = DEFAULT_BODY_PLACEHOLDER) -> bytes:
    """Build a minimal bank template: a cover line, the placeholder, a footer line."""
    doc = Document()
    doc.add_heading("\u3007\u3007\u9280\u884c \u5be9\u67fb\u90e8", level=0)  # cover/branding
    doc.add_paragraph("COVER-MARKER")
    doc.add_paragraph(placeholder)
    doc.add_paragraph("FOOTER-MARKER")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _all_text(docx_bytes: bytes) -> str:
    """Concatenate all paragraph + table-cell text from a .docx for assertions."""
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class TestTemplateInjection:
    def test_figures_preserved_verbatim(self) -> None:
        out = build_keikakusho_docx_from_template(_DRAFT, _template_bytes())
        text = _all_text(out)
        for figure in _FIGURES:
            assert figure in text, f"figure {figure} must survive verbatim"

    def test_template_surroundings_preserved(self) -> None:
        out = build_keikakusho_docx_from_template(_DRAFT, _template_bytes())
        text = _all_text(out)
        assert "COVER-MARKER" in text
        assert "FOOTER-MARKER" in text
        assert "\u3007\u3007\u9280\u884c \u5be9\u67fb\u90e8" in text

    def test_placeholder_is_removed(self) -> None:
        out = build_keikakusho_docx_from_template(_DRAFT, _template_bytes())
        assert DEFAULT_BODY_PLACEHOLDER not in _all_text(out)

    def test_body_lands_between_cover_and_footer(self) -> None:
        """The injected body must sit where the placeholder was (not at the end)."""
        out = build_keikakusho_docx_from_template(_DRAFT, _template_bytes())
        doc = Document(io.BytesIO(out))
        texts = [p.text for p in doc.paragraphs]
        cover_idx = next(i for i, t in enumerate(texts) if t == "COVER-MARKER")
        footer_idx = next(i for i, t in enumerate(texts) if t == "FOOTER-MARKER")
        body_idx = next(i for i, t in enumerate(texts) if "\u00a5123,456,789" in t)
        assert cover_idx < body_idx < footer_idx

    def test_tables_render_as_word_tables(self) -> None:
        out = build_keikakusho_docx_from_template(_DRAFT, _template_bytes())
        doc = Document(io.BytesIO(out))
        assert doc.tables, "the Markdown table must become a real Word table"

    def test_missing_placeholder_raises(self) -> None:
        bad = _template_bytes(placeholder="NO-PLACEHOLDER-HERE")
        with pytest.raises(TemplateError):
            build_keikakusho_docx_from_template(_DRAFT, bad)

    def test_custom_placeholder(self) -> None:
        tpl = _template_bytes(placeholder="<<BODY>>")
        out = build_keikakusho_docx_from_template(_DRAFT, tpl, placeholder="<<BODY>>")
        assert "\u00a5123,456,789" in _all_text(out)
        assert "<<BODY>>" not in _all_text(out)


class TestDefaultAndSettings:
    def test_default_mode_unchanged(self) -> None:
        """With no template, the bare document still contains every figure."""
        out = build_keikakusho_docx(_DRAFT)
        text = _all_text(out)
        for figure in _FIGURES:
            assert figure in text
        # No template surroundings leak into the default export.
        assert "COVER-MARKER" not in text

    def test_settings_without_template_falls_back_to_default(self) -> None:
        class _S:
            keikakusho_docx_template = ""

        out = build_keikakusho_docx_for_settings(_DRAFT, _S())
        assert "\u00a5123,456,789" in _all_text(out)
        assert "COVER-MARKER" not in _all_text(out)

    def test_settings_with_template_uses_it(self, tmp_path: object) -> None:
        from pathlib import Path

        template_file = Path(str(tmp_path)) / "bank_template.docx"
        template_file.write_bytes(_template_bytes())

        class _S:
            keikakusho_docx_template = str(template_file)

        out = build_keikakusho_docx_for_settings(_DRAFT, _S())
        text = _all_text(out)
        assert "COVER-MARKER" in text  # template was used
        assert "\u00a5123,456,789" in text  # body injected, figure preserved

    def test_settings_with_missing_template_path_falls_back(self) -> None:
        class _S:
            keikakusho_docx_template = "/nonexistent/path/bank_template.docx"

        out = build_keikakusho_docx_for_settings(_DRAFT, _S())
        # Falls back to default rather than raising.
        assert "\u00a5123,456,789" in _all_text(out)
        assert "COVER-MARKER" not in _all_text(out)
