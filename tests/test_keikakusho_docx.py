"""Tests for the deterministic Keikakusho → DOCX exporter.

The load-bearing invariant is numeric preservation: converting the draft to
Word must never add, drop, or alter a yen figure. We assert this by reusing the
project's own ``extract_yen_values`` over the source Markdown and over the text
extracted back out of the generated ``.docx``.
"""

from __future__ import annotations

import io
from collections import Counter

from app.backend.analysis.numeric_preservation import extract_yen_values
from app.backend.export.keikakusho_docx import build_keikakusho_docx, docx_filename
from docx import Document

_DRAFT = "\n".join(
    [
        "# 経営改善計画書（Keiei Kaizen Keikakusho）",
        "",
        "- 企業名（Company）: テスト製造株式会社",
        "- 法人番号（Hojin Bango）: 1234567890123",
        "- 債務者区分（FSA classification）: 要注意先",
        "",
        "## 1. 現状分析（Current position）",
        "",
        "- 売上（Uriage）: ¥150,000,000",
        "- 売上原価（Uriage Genka）: ¥120,000,000",
        "- 経常利益（Keijo Rieki）: -¥5,000,000",
        "- 資金繰りギャップ（Working-capital gap）: -¥86,112,067",
        "",
        "## 2. 改善施策（Turnaround strategy）",
        "",
        "### 価格転嫁の実行",
        "",
        "- 期待される経常利益改善: ¥54,000,000 / 年",
        "",
        "## 3. 実行計画（Action plan）",
        "",
        "1. 施策の実行体制を構築する。",
        "2. 月次で進捗をモニタリングする。",
        "",
    ]
)


def _docx_text(data: bytes) -> str:
    """Extract all paragraph text from generated ``.docx`` bytes."""
    document = Document(io.BytesIO(data))
    return "\n".join(p.text for p in document.paragraphs)


def _docx_all_text(data: bytes) -> str:
    """Extract paragraph AND table-cell text from generated ``.docx`` bytes.

    ``python-docx`` exposes table cells separately from ``document.paragraphs``,
    so a numeric-preservation check that ignores tables would miss any figure
    moved into a Word table. This walks both.
    """
    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


_TABLE_DRAFT = "\n".join(
    [
        "## 4. 損益計画（Recovery projection）",
        "",
        "| 月 (Month) | 月次改善額 (Uplift) | 経常利益 (Keijo Rieki) | EWS |",
        "| ---: | ---: | ---: | ---: |",
        "| 1 | ¥3,000,000 | -¥2,000,000 | 58.20 |",
        "| 2 | ¥3,000,000 | ¥1,000,000 | 41.10 |",
        "",
    ]
)


def test_markdown_table_becomes_a_word_table() -> None:
    """A Markdown table renders as a real Word table, not pipe-delimited text."""
    document = Document(io.BytesIO(build_keikakusho_docx(_TABLE_DRAFT)))
    assert len(document.tables) == 1
    table = document.tables[0]
    # Header + 2 body rows = 3 rows; 4 columns.
    assert len(table.rows) == 3
    assert len(table.columns) == 4
    # Header cells copied verbatim.
    header = [c.text for c in table.rows[0].cells]
    assert header[0] == "月 (Month)"
    assert header[3] == "EWS"
    # A body figure is carried verbatim into its cell.
    assert table.rows[1].cells[1].text == "¥3,000,000"


def test_table_delimiter_row_is_dropped() -> None:
    """The ``| --- |`` layout row never leaks into the document text."""
    all_text = _docx_all_text(build_keikakusho_docx(_TABLE_DRAFT))
    assert "---" not in all_text
    # And no raw pipe-delimited row survived as a paragraph.
    assert "| 月" not in all_text


def test_table_preserves_every_yen_figure() -> None:
    """Yen figures inside a table survive byte-for-byte (paragraphs + cells)."""
    data = build_keikakusho_docx(_TABLE_DRAFT)
    source = Counter(extract_yen_values(_TABLE_DRAFT))
    rendered = Counter(extract_yen_values(_docx_all_text(data)))
    assert rendered == source
    assert source  # guard: the draft actually contains yen figures


def test_no_table_draft_creates_no_tables() -> None:
    """A draft without a table produces no Word tables (behaviour unchanged)."""
    document = Document(io.BytesIO(build_keikakusho_docx(_DRAFT)))
    assert document.tables == []


def test_build_keikakusho_docx_returns_nonempty_bytes() -> None:
    data = build_keikakusho_docx(_DRAFT)
    assert isinstance(data, bytes)
    assert len(data) > 0
    # A .docx is a zip archive — starts with the PK signature.
    assert data[:2] == b"PK"


def test_docx_preserves_every_yen_figure() -> None:
    """The DOCX must carry exactly the same multiset of yen values as the draft."""
    data = build_keikakusho_docx(_DRAFT)
    source = Counter(extract_yen_values(_DRAFT))
    rendered = Counter(extract_yen_values(_docx_text(data)))
    assert rendered == source


def test_docx_carries_classification_and_company_text() -> None:
    text = _docx_text(build_keikakusho_docx(_DRAFT))
    assert "要注意先" in text
    assert "テスト製造株式会社" in text


def test_docx_filename_is_safe() -> None:
    assert docx_filename("テスト 製造") == "keikakusho_テスト_製造.docx"
    assert docx_filename("") == "keikakusho_keikakusho.docx"
    assert docx_filename("a/b\\c") == "keikakusho_a_b_c.docx"
