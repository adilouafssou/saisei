"""Shared, deterministic Markdown → DOCX (Word) renderer.

Japanese banks and FSA examiners exchange regulated documents as Word, not
Markdown. Several deterministic artifacts in this package are produced as
Markdown first (the Keikakusho draft, the Feature 7 explainability report); this
module is the ONE renderer that turns any such Markdown into an editable
``.docx`` so each artifact gets a Word path without duplicating the converter.

Numeric-preservation invariant
------------------------------
The one inviolable project rule is that no step may add, drop, or alter a
number. This converter upholds it structurally: it walks the Markdown
**line-by-line** and copies each line's text **verbatim** into a Word paragraph
(or table cell), mapping only the Markdown *structure* (``#`` heading levels,
``-`` / numbered list items, ``| ... |`` tables) to Word styling. It never
parses, reformats, or re-renders a figure — the bytes of every yen value are
carried across unchanged.

The function is pure and deterministic: same Markdown in → same ``.docx`` bytes
out (no network, no LLM). ``python-docx`` is the only dependency (already a
project dependency; no new one is added).
"""

from __future__ import annotations

import io
import re
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Pt

if TYPE_CHECKING:
    from docx.document import Document as DocumentT

__all__ = ["render_markdown_to_docx", "render_markdown_into_document"]

#: A Markdown ATX heading: leading ``#``\ s, a space, then the text.
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
#: An unordered list item: ``- `` (or ``* ``) then the text.
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
#: An ordered list item: ``1. `` then the text (number kept verbatim in text).
_ORDERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")
#: A Markdown table row: a line that starts and ends with a pipe (after strip).
_TABLE_ROW_RE = re.compile(r"^\|.*\|$")
#: A Markdown table delimiter row: only pipes, dashes, colons, and spaces
#: (e.g. ``| --- | ---: |``). Dropped from the rendered table (it is layout,
#: not content), so it never leaks into the document as a junk line.
_TABLE_DELIM_RE = re.compile(r"^\|[\s:|-]+\|$")
#: Heading levels above Word's built-in ``Heading 1..4`` collapse to 4.
_MAX_WORD_HEADING = 4


def _strip_inline_emphasis(text: str) -> str:
    """Remove Markdown ``**bold**`` / ``*italic*`` markers, keeping the text.

    Only the surrounding emphasis markers are removed; digits, currency markers
    (¥ / 円), commas, and signs are never touched, so yen figures survive
    byte-for-byte. Bold/italic styling is intentionally dropped (not
    reconstructed) to keep the converter trivial and number-safe.
    """
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    return text


def _split_table_cells(row: str) -> list[str]:
    """Split a ``| a | b |`` Markdown row into verbatim, trimmed cell strings.

    Only the outer pipes and the inter-cell pipes are removed and each cell is
    surface-trimmed; digits, currency markers, commas, and signs inside a cell
    are never touched, so yen figures survive byte-for-byte into the Word table.
    """
    inner = row.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [_strip_inline_emphasis(cell.strip()) for cell in inner.split("|")]


def _flush_table(document: DocumentT, table_rows: list[str]) -> None:
    """Render a buffered block of Markdown table rows as a real Word table.

    ``table_rows`` are the raw ``| ... |`` lines in order (delimiter rows already
    excluded). The first row is treated as the header. Every cell is copied
    verbatim via :func:`_split_table_cells`, so the numeric-preservation
    invariant holds inside the table exactly as it does for paragraphs. Ragged
    rows are tolerated by sizing the table to the widest row and leaving missing
    cells blank. A single-row (header-only) block still renders as a 1-row table.
    """
    if not table_rows:
        return
    parsed = [_split_table_cells(r) for r in table_rows]
    n_cols = max(len(cells) for cells in parsed)
    table = document.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"
    for cells in parsed:
        row_cells = table.add_row().cells
        for i in range(n_cols):
            row_cells[i].text = cells[i] if i < len(cells) else ""


def render_markdown_into_document(document: DocumentT, markdown: str) -> None:
    """Render deterministic Markdown into an EXISTING Word document, in place.

    Identical line-by-line, verbatim, number-safe conversion as
    :func:`render_markdown_to_docx`, but appends the resulting paragraphs and
    tables to a caller-provided ``document`` instead of a fresh one. This is what
    lets the Keikakusho body be injected into a bank's house template (cover
    page, letterhead, styles) without duplicating the converter or touching a
    single figure.

    The numeric-preservation invariant is unchanged: each line's text is copied
    verbatim; only Markdown *structure* maps to Word styling.

    Args:
        document: The target ``python-docx`` Document to append into.
        markdown: The deterministic Markdown source of truth.
    """
    # Buffer consecutive table rows so a contiguous block becomes ONE Word table.
    table_buffer: list[str] = []

    def _flush() -> None:
        _flush_table(document, table_buffer)
        table_buffer.clear()

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        # --- Table handling (buffer contiguous rows; drop the delimiter). ---
        if _TABLE_ROW_RE.match(stripped):
            if not _TABLE_DELIM_RE.match(stripped):
                table_buffer.append(stripped)
            continue
        # Any non-table line ends an open table block.
        if table_buffer:
            _flush()

        if not stripped:
            # Blank line → a blank paragraph preserves vertical rhythm.
            document.add_paragraph("")
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            level = min(len(heading.group(1)), _MAX_WORD_HEADING)
            document.add_heading(_strip_inline_emphasis(heading.group(2)), level=level)
            continue

        bullet = _BULLET_RE.match(stripped)
        if bullet:
            document.add_paragraph(_strip_inline_emphasis(bullet.group(1)), style="List Bullet")
            continue

        ordered = _ORDERED_RE.match(stripped)
        if ordered:
            # Keep the author's own numbering verbatim in the text (do not let
            # Word renumber), so any figure in the line is untouched.
            document.add_paragraph(
                f"{ordered.group(1)}. {_strip_inline_emphasis(ordered.group(2))}"
            )
            continue

        document.add_paragraph(_strip_inline_emphasis(stripped))

    # Flush a table that runs to the very end of the document.
    if table_buffer:
        _flush()


def render_markdown_to_docx(markdown: str) -> bytes:
    """Render a deterministic Markdown document to ``.docx`` bytes.

    Walks ``markdown`` line-by-line and copies each line's text verbatim into a
    Word paragraph, mapping only the Markdown *structure* (heading levels, list
    items, and tables) to Word styling. No figure is ever reformatted.

    Markdown tables (a header row, a ``| --- |`` delimiter, then body rows) are
    rendered as real Word tables with every cell copied verbatim; the delimiter
    row is layout and is dropped.

    Pure and deterministic: same Markdown in → same ``.docx`` structure out.

    Args:
        markdown: The deterministic Markdown source of truth.

    Returns:
        The generated ``.docx`` file as bytes.
    """
    document = Document()

    # Use a CJK-friendly base font so the document renders cleanly in Word even
    # if the reviewer's default font lacks Japanese glyphs. This styles the
    # document; it does not alter any text content.
    normal = document.styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal.font.size = Pt(11)

    render_markdown_into_document(document, markdown)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
